"""
生成青付安产品方案书 Word 文档 V3 — 含三大用户画像研究 + 修复表格格式
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
import os

doc = Document()

# ── 页面设置 ──
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

# ── 全局样式 ──
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = '微软雅黑'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    hs.font.color.rgb = RGBColor(0x24, 0x59, 0x57)
    hs.font.bold = True
    if i == 1:
        hs.font.size = Pt(22)
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(12)
    elif i == 2:
        hs.font.size = Pt(16)
        hs.paragraph_format.space_before = Pt(20)
        hs.paragraph_format.space_after = Pt(8)
    else:
        hs.font.size = Pt(13)
        hs.paragraph_format.space_before = Pt(14)
        hs.paragraph_format.space_after = Pt(6)

def add_para(text, bold=False, size=10.5, color=None, align=None, space_after=6, indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_code(text, size=9):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F5F5F5')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def set_cell_shading(cell, color_hex):
    """给单元格设置背景色"""
    tcPr = cell._element.get_or_add_tcPr()
    # 清除已有的 shading
    for existing in tcPr.findall(qn('w:shd')):
        tcPr.remove(existing)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, size=9, color=None, align='left', font_name='微软雅黑'):
    """设置单元格文本和格式"""
    # 清除默认空段落
    for p in cell.paragraphs:
        p.clear()
    p = cell.paragraphs[0]
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 段落间距
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

def add_table(headers, rows, col_widths=None, header_color='48A9A6'):
    """创建带薄荷绿表头和斑马纹的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 去掉默认表格样式，手动控制格式
    table.style = 'Table Grid'

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_color)
        set_cell_text(cell, h, bold=True, size=9, color=(255,255,255), align='center')

    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            if r % 2 == 1:  # 斑马纹（偶数行）
                set_cell_shading(cell, 'F5FAFA')
            set_cell_text(cell, val, size=9)

    # 列宽
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # 表后空行
    return table

# ═══════════════════════════════════════
# 封面
# ═══════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
add_para('青付安（QingFuAn）', bold=True, size=28, color=(0x24,0x59,0x57), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('产品方案书', bold=True, size=20, color=(0x48,0xA9,0xA6), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
add_para('青年预付消费管理 PWA', size=12, color=(0x63,0x8F,0x8D), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
add_para('版本：V3.0  |  更新日期：2026-08-11', size=10, color=(0x63,0x8F,0x8D), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
add_para('在线体验：https://lin-hmpp.github.io/qingfuan/', size=9, color=(0x4A,0x7A,0x77), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('项目仓库：https://github.com/Lin-HMPP/qingfuan', size=9, color=(0x4A,0x7A,0x77), align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ═══════════════════════════════════════
# 目录
# ═══════════════════════════════════════
doc.add_heading('目录', level=1)
toc_items = [
    '一、产品概述',
    '二、用户画像研究（基于 87 份问卷 + 9 份访谈）',
    '三、功能架构',
    '四、核心功能设计',
    '五、风险规则引擎',
    '六、数据模型',
    '七、技术方案',
    '八、UI 设计规范',
    '九、安全与隐私',
    '十、产品截图索引',
    '十一、已知限制与后续方向',
    '附录',
]
for item in toc_items:
    add_para(item, size=11, space_after=4)

doc.add_page_break()

# ═══════════════════════════════════════
# 一、产品概述
# ═══════════════════════════════════════
doc.add_heading('一、产品概述', level=1)

doc.add_heading('1.1 产品定义', level=2)
add_para('青付安是一款面向青年消费者的预付消费管理 PWA。覆盖办卡前风险评估、办卡后资产追踪、纠纷时证据归集三个环节。核心设计理念：不是劝用户"别买"，而是帮用户看清风险后自己做决定。')

doc.add_heading('1.2 产品形态', level=2)
add_table(
    ['维度', '说明'],
    [
        ['形态', 'PWA，浏览器打开即用，可添加至主屏幕'],
        ['平台', '移动端，适配微信内置浏览器'],
        ['技术', 'Vue 3 + Vite 5，纯前端，localStorage 全量本地存储，零服务端'],
        ['预设场景', '健身/舞蹈、培训课程、摄影套餐、美容美发（支持自定义）'],
        ['公开地址', 'https://lin-hmpp.github.io/qingfuan/'],
    ],
    [3, 11]
)

doc.add_heading('1.3 解决的核心问题', level=2)
add_table(
    ['用户痛点', '青付安解法'],
    [
        ['办卡前高估使用频率，用"理想计划"代替"真实行为"', '套餐录入 → 频率-有效期匹配分析 → 多情景成本测算'],
        ['办卡后忘记剩余次数、到期时间、多卡分散难管理', '资产列表统一视图 + 持仓卡实时权益计算 + 到期预警'],
        ['核销记录散落各处，权益算不清', '统一核销打卡，无限次模式自动追踪累计到店与回本'],
        ['门店纠纷/跑路时缺证据', '8 类材料清单 + 一键打包 HTML 维权报告'],
    ],
    [4, 10]
)

# ═══════════════════════════════════════
# 二、用户画像研究
# ═══════════════════════════════════════
doc.add_heading('二、用户画像研究', level=1)
add_para('本章基于 87 份问卷（有效购买者 52 人）+ 9 份深度访谈，归纳出三类行为画像。三类画像是"行为原型"而非互斥的人群标签——同一用户会随消费金额、商户规模、需求稳定性和受损经历在三种模式间切换。', size=9.5, color=(0x63,0x8F,0x8D))

doc.add_heading('2.1 研究样本', level=2)
add_table(
    ['指标', '数据'],
    [
        ['访谈材料', '9 份（8 份完整文字访谈，1 份含图片记录）'],
        ['问卷总数', '87 份，其中 52 人为近两年预付消费有效购买者'],
        ['有效购买者构成', '69.2% 在校本科生/专科生；65.4% 月可支配 1,000-2,999 元'],
        ['高频消费类型', '软件影音会员、餐饮储值、健身舞蹈、美容服务、教育课程'],
    ],
    [5, 9]
)

doc.add_heading('2.2 三类画像总览', level=2)
add_table(
    ['维度', '画像一：未来自我下注者', '画像二：刚需精算执行者', '画像三：风险警觉守门人'],
    [
        ['需求确定性', '愿望强，未来稳定性低', '需求明确且已被日常行为验证', '需求可高可低，先看风险是否可控'],
        ['决策核心', '"交了钱就会坚持"', '"确定会用，长期买更便宜"', '"靠不靠谱，能不能退出"'],
        ['最大误区', '高估未来时间和自律', '低估多卡管理成本和隐性条款', '过度依赖个人经验，只防已见过的坑'],
        ['主要损失', '闲置、遗忘、单次成本飙升', '卡项分散、规则遗漏、自动续费', '退款困难、主体不明、证据不足'],
        ['最需要的帮助', '付款前冷静期 + 现实频率测算', '快速比较、低摩擦记录、统一管理', '条款识别、主体核验、证据链与行动指引'],
    ],
    [2.5, 3.5, 3.8, 4.2],
)

doc.add_heading('2.3 画像一：未来自我下注者', level=2)
add_para('"我不是在买卡，我是在买一个会坚持的自己"', bold=True, size=10, color=(0x48,0xA9,0xA6))

add_para('典型轮廓', bold=True, size=10)
add_para('18-25 岁学生或初入职场者。时间安排经常变化——学业、实习、考试、假期都会影响使用频率。常见于健身年卡、舞蹈/兴趣课程、美甲多次卡，以及带有"变美、变瘦、提升技能"承诺的套餐。', indent=0.5)

add_para('核心心理机制', bold=True, size=10)
add_para('预付款在此类用户身上同时承担三种心理作用：① 承诺装置——用沉没成本推动未来行动；② 身份确认——办卡本身让人短暂感觉"我已经开始改变"；③ 愿望放大——销售话术、朋友同行、赠送月份共同把"想试试"包装成"我会长期坚持"。', indent=0.5)
add_para('最深层矛盾：不是不知道年卡可能浪费，而是把"浪费风险"也当成自律压力的一部分。单纯告诉她"年卡不划算"不足以阻止购买。', indent=0.5)

add_para('决策路径', bold=True, size=10)
add_para('愿望触发（想减肥/学技能/充实生活）→ 出现优惠（学生价/赠月/限时）→ 以"我应该每周去三次"替代"过去三个月我实际去了几次"→ 长期方案看似月均最低 → 短期热情 → 现实摩擦（课忙/距离/假期/预约难）→ 闲置遗忘', indent=0.5)

add_para('问卷数据支撑', bold=True, size=10)
add_para('• 52 名购买者中 50.0% 的最近一次消费已出现低频、长期未用或到期未用完', indent=0.5)
add_para('• 上述低频/闲置者中 100% 从未认真做过单次成本测算', indent=0.5)
add_para('• 73.1% 遇到过忘记使用或到期仍有余额/次数', indent=0.5)
add_para('• 即使是在低频/闲置者中，仍有 53.8% 当初认为自己存在"稳定需求"', indent=0.5)

add_para('青付安的设计回应', bold=True, size=10)
add_para('• 付款前多情景测算（乐观/现实/保守三种使用频率下的真实成本）', indent=0.5)
add_para('• 套餐录入页的"频率智能匹配"——用过去行为数据替代理想计划', indent=0.5)
add_para('• 决策卡不简单给"买/不买"建议，而是说"如果你每周能去 X 次就划算，否则按次消费更划算"', indent=0.5)
add_para('• 新增同类卡时提醒"你仍有一张同类卡未用完"', indent=0.5)

doc.add_heading('2.4 画像二：刚需精算执行者', level=2)
add_para('"我已经在稳定使用，所以长期买才是真的省"', bold=True, size=10, color=(0x48,0xA9,0xA6))

add_para('典型轮廓', bold=True, size=10)
add_para('有明确日常习惯的学生、研究生或年轻职场人。不同于画像一——不是因为长期卡便宜才制造需求，而是先有稳定需求（如每周固定健身 4 次），再用长期卡降低单位成本。', indent=0.5)

add_para('决策路径', bold=True, size=10)
add_para('需求先存在 → 按距离/品牌/评分缩小范围 → 体验或验证（线下看/体验卡/看评价）→ 比较月卡/季卡/年卡成本 → 匹配现实周期（学期/假期/工作量）→ 持续使用', indent=0.5)

add_para('问卷数据支撑', bold=True, size=10)
add_para('• 最近一次消费正常使用/结束者占 46.2%。其中：70.8% 以"长期稳定需求"为购买原因；75.0% 同时受单价更低驱动；70.8% 做过成本估算；79.2% 会用工具查询记录', indent=0.5)
add_para('• 但即使在这组人中：48.1% 有至少 2 个未用完套餐；88.0% 不能准确说清全部剩余权益和到期时间；76.0% 遇到过忘记使用或到期未用完', indent=0.5)

add_para('被忽略的结构性痛点', bold=True, size=10)
add_para('• 多卡分散在商户 App、微信聊天、相册和纸质合同里', indent=0.5)
add_para('• 自动续费失察——单个金额小，长期累积才发现', indent=0.5)
add_para('• 规则只看与当下有关的部分——退款、暂停、迁店条款到需要时才发现限制', indent=0.5)
add_para('• 管理意愿低于管理需求——单卡时觉得不用管，卡多了才感到混乱', indent=0.5)

add_para('青付安的设计回应', bold=True, size=10)
add_para('• 快速录入（5 字段、~10 秒）——降低"不值得花时间管"的心理门槛', indent=0.5)
add_para('• 资产列表统一视图——多卡剩余权益一目了然', indent=0.5)
add_para('• 只轻量提醒——偏离计划或临近关键节点时才通知，不制造打扰感', indent=0.5)
add_para('• 解释计算依据——规则引擎每条结论都附计算过程，满足"追问如何得出的"需求', indent=0.5)

doc.add_heading('2.5 画像三：风险警觉守门人', level=2)
add_para('"便宜不够，我要先确认钱交给谁、出事能不能退"', bold=True, size=10, color=(0x48,0xA9,0xA6))

add_para('典型轮廓', bold=True, size=10)
add_para('自己或身边人经历过健身房跑路、强制推销、转卡扣费、退款拖延或口头承诺不兑现。常见于高金额健身/培训卡、小型门店、美发推销、主体关系复杂的消费。决策前提不是"绝对没风险"，而是风险能被看见、被证明、有退出路径。', indent=0.5)

add_para('四大风险敏感区', bold=True, size=10)
add_para('① 主体风险：合同方、门店、品牌方、收款方是否一致，谁真正承担责任', indent=0.5)
add_para('② 退出风险：退款、转卡、暂停、延期是否可行，手续费或扣课规则', indent=0.5)
add_para('③ 履约风险：商户停业、迁店、人员更换、服务缩水后的应对', indent=0.5)
add_para('④ 证据风险：口头承诺是否写入合同，聊天、宣传、付款记录能否在出事后找回', indent=0.5)

add_para('问卷数据支撑', bold=True, size=10)
add_para('• 55.8% 的购买者没有核对合同名称、门店经营者与收款方是否对应', indent=0.5)
add_para('• 仅 15.4% 收到协议后认真看了主要内容；23.1% 根本没有收到合同或协议', indent=0.5)
add_para('• 明确遇到退款/履约/宣传不一致/停业等风险事件的用户约 19.2%', indent=0.5)
add_para('• 在这组风险经历者中：60.0% 材料保存能力较弱；70.0% 明确需要付款前规则检查', indent=0.5)

add_para('信任设计要求', bold=True, size=10)
add_para('对这类用户，产品本身也会被当作风险对象审视。青付安的设计原则：① 明确说明数据存在本地、不上传；② 规则结论区分"事实提取"和"风险提示"，不夸大确定性；③ 展示判断依据——具体哪条规则命中、缺失了什么信息；④ 支持本地保存、导出和删除。', indent=0.5)

add_para('青付安的设计回应', bold=True, size=10)
add_para('• R7 主体一致性核验——自动比对合同方、门店、收款方是否为同一实体', indent=0.5)
add_para('• R17 平台团购支付风险——识别非直接付商家的维权链路风险', indent=0.5)
add_para('• 证据资料夹 8 类材料清单 + 缺失可点击上传——变被动"事后找材料"为主动"事前归集"', indent=0.5)
add_para('• 一键打包 HTML 维权报告——出问题时材料已经在手', indent=0.5)

doc.add_heading('2.6 画像间的深层关系', level=2)
add_para('用户不是固定属于某一类。最常见的迁移路径：未来自我下注者 → 发生闲置或受损 → 风险警觉守门人。但迁移未必完整——用户可能学会防商家，却仍然高估自己的使用频率。', indent=0.5)
add_para('因此产品需要区分两类风险：① 自我履约风险——用户自己是否有时间、有习惯、有持续需求；② 商户履约风险——商户是否透明、稳定、能够退款和持续服务。这也是青付安同时做"使用频率分析"和"合同规则检查"两件事的底层逻辑。', indent=0.5)
add_para('"稳定需求"是最容易被误判的自我陈述。问卷中 63.5% 购买者以"长期稳定需求"为购买原因，但 50.0% 最近一次消费已出现低频/闲置/未用完。甚至在低频闲置者中，仍有 53.8% 当初认为自己存在稳定需求。因此青付安的规则引擎不依赖用户自述的"使用频率"，而是要求填写计划频率后与之做算法比对（R2）。', indent=0.5)

doc.add_page_break()

# ═══════════════════════════════════════
# 三、功能架构
# ═══════════════════════════════════════
doc.add_heading('三、功能架构', level=1)

doc.add_heading('3.1 页面结构（11 个页面 + 4 Tab）', level=2)

pages = [
    ('🏠 首页（/home）', [
        '到期提醒卡片（30 天内到期置顶）',
        '双卡片入口（"购买前先检查"/"我的预付资产"）',
        '预付总额 + 卡数统计 + 证据资料夹快捷入口',
        '场景标签（4 预设 + 自定义）+ ⊕ 快速录入',
    ]),
    ('⚡ 快速录入（/quick-input）—— 服务画像二"低摩擦记录"', [
        '5 字段极简建卡：场景标签 + 门店名 + 总价 + 次卡/充卡切换 + 期限',
        '~10 秒完成，直接创建资产卡 + 自动创建同名资料夹',
    ]),
    ('📝 套餐录入（/package-input）—— 服务画像一"付款前冷静测算"', [
        '模块一：基础费用 + 实时成本预览（乐观/现实/保守三种估算）',
        '模块二：使用规划 + 频率智能匹配（比对计划与实际，服务 R2 规则）',
        '模块三：签约&收款主体（5 种支付渠道，服务 R7/R17）',
        '模块四：履约规则选择（退款/转卡/暂停，场景化预设选项）',
        '模块五：促销附加说明（选填）',
        '8 类材料上传 + 每 10 秒自动保存草稿',
    ]),
    ('📊 决策卡（/decision-card）—— 17 条规则结论可视化', [
        '结论横幅（绿/橙/红三档）+ 套餐速览',
        '花费算账：单次成本/月预算对比/消耗节奏/回本参考',
        '需关注的问题（高/中风险分类）+ 已通过检查（可折叠）',
        '付款前行动清单（去重）+ 内嵌六维度风险详情',
    ]),
    ('📋 风险报告（/risk-report，可选）', [
        '六维度逐板块展示 + 双情景成本测算（普通模式）/ 充卡价值分析（无限次模式）',
    ]),
    ('💳 资产列表 + 🏷️ 持仓卡 + ✓ 核销录入', [
        '资产统一管理视图 + 实时权益计算 + 无限次到店打卡 + 暂停/过期拦截',
    ]),
    ('📁 证据资料夹（/evidence-folder）—— 服务画像三"证据保险箱"', [
        '文件夹列表 → 点击进入：其余隐藏，仅显示选中卡片（高亮）+ 详情',
        '8 类材料完整性检查清单（缺失项可点击直达上传）',
        '一键打包 HTML 维权报告（图片内嵌，可直接打印）',
    ]),
    ('👤 我的（/mine）', [
        '三栏统计 + PIN 锁定/解锁 + 17 条规则说明 + 数据管理',
    ]),
]

for title, items in pages:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = RGBColor(0x24, 0x59, 0x57)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    for item in items:
        add_para('    ' + item, size=9.5, space_after=1)

doc.add_heading('3.2 公共组件（11 个）', level=2)
add_table(
    ['组件', '用途'],
    [
        ['tab-bar', '底部四 Tab（首页/资产/证据/我的），内联 SVG 响应式变色'],
        ['toast', '全局深色 Toast，3 秒自动消失，DOM 兜底'],
        ['pin-lock', 'PIN 数字键盘（设置/验证模式）'],
        ['scene-picker', '场景选择弹窗（4 预设 + 自定义入口）'],
        ['image-picker', '图片采集（拍摄/相册/文件 → base64，支持多选）'],
        ['refund-checklist', '退款材料核对清单（支持无限次时间比例模式）'],
        ['pause-transfer', '暂停锁卡 + 编辑套餐有效期'],
        ['asset-confirm', '确认生成资产弹窗'],
        ['exit-confirm', '退出确认弹窗'],
        ['write-off-detail', '核销详情弹窗（编辑/删除）'],
        ['package-loading', '打包导出加载动画'],
    ],
    [4, 10]
)

# ═══════════════════════════════════════
# 四、核心功能设计
# ═══════════════════════════════════════
doc.add_heading('四、核心功能设计', level=1)

doc.add_heading('4.1 场景化文案', level=2)
add_para('4 种预设场景各有独立配置（SCENE_COPY），覆盖输入框 placeholder、规则选项术语、核销单位标签。例如培训场景用"休学"替代"暂停"，核销单位用"课时"替代"次"。切换场景时文案动态替换，降低用户理解成本。')

doc.add_heading('4.2 双录入路径', level=2)
add_table(
    ['路径', '字段数', '耗时', '适用场景', '对应画像'],
    [
        ['快速录入', '5', '~10 秒', '已办卡、仅需记账', '画像二（低摩擦）'],
        ['套餐录入', '5 模块 + 8 类材料', '~3 分钟', '办卡前、需风险评估', '画像一 + 画像三'],
    ],
    [2.5, 1.5, 1.5, 4, 4.5],
)

doc.add_heading('4.3 无限次模式（充卡/年卡）', level=2)
add_para('传统次卡逻辑对充年卡不适用。无限次模式的全套适配：持仓卡展示累计到店 + 日均成本 + 时间进度条；核销页自动切换"到店打卡"模式；风险报告展示回本所需最低到店频率；退款清单按时间比例计算。')

doc.add_heading('4.4 证据资料夹闭环', level=2)
add_para('对标画像三"证据保险箱"需求。从录入上传 → 自动归入对应资产文件夹 → 8 类完整性检查（缺什么直接点）→ 一键打包 HTML 维权报告。报告内嵌 base64 图片，可直接打印提交。')

doc.add_heading('4.5 多情景成本测算', level=2)
add_para('对标画像一的核心诉求——不只展示"每次只要 XX 元"，而是给出三种情景：① 理想频率下的单次价格；② 按过去真实行为推算的价格；③ 考虑假期/考试/兴趣衰减后的保守价格。决策卡的"花费算账"模块直接呈现这些计算结果。')

# ═══════════════════════════════════════
# 五、风险规则引擎
# ═══════════════════════════════════════
doc.add_heading('五、风险规则引擎', level=1)

doc.add_heading('5.1 规则清单（17 条）', level=2)
add_table(
    ['编号', '规则', '触发条件', '等级'],
    [
        ['R1', '单次/日均成本核算', '成本 > 市场参考价 2×', '中'],
        ['R2', '有效期-频率匹配', '所需频率与计划频率差距 >1.5', '高'],
        ['R3', '合同/凭证可得性', '无书面合同', '中'],
        ['R4', '退款条款清晰度', '仅口头承诺', '高'],
        ['R5', '转卡/暂停/延期条款', '规则不明确', '中'],
        ['R6', '迁址/停业应对条款', '长周期 + 无条款', '中'],
        ['R7', '主体一致性核验', '店名≠合同≠收款方，或个人收款', '高'],
        ['R8', '高金额预付', '总价 > 月预算 3 倍', '高'],
        ['R9', '赠品/优惠限制', '有赠品但规则不清', '中'],
        ['R10', '到期预警', '剩余 ≤30 天', '高'],
        ['R11', '使用频率异常', '实际/计划 < 0.5', '高'],
        ['R12', '材料留存完整性', '缺失 ≥3 类', '高'],
        ['R13', '退款前置条件', '退款规则未知', '中'],
        ['R14', '服务变更/价格调整', '默认跟踪', '低'],
        ['R15', '场景专属检测', '默认跟踪', '低'],
        ['R16', '退款渠道核验', '商家通常不主动告知', '中'],
        ['R17', '平台团购支付风险', '非直接付商家（维权链路更长）', '中'],
    ],
    [1.2, 4, 5.8, 1]
)

doc.add_heading('5.2 六维度评分', level=2)
add_para('17 条规则归入 6 个维度，每个维度 1-5 分。综合评定：高风险项 ≥3 → 高风险；高风险项 ≥1 或中风险 ≥4 → 中风险；其他 → 低风险。')
add_table(
    ['维度', '关联规则', '对应画像痛点'],
    [
        ['预付压力', 'R1, R2, R8', '画像一的"理想频率替代真实频率"'],
        ['履约时限', 'R9, R10, R11', '画像一的"忘记使用/到期仍有余额"'],
        ['合约权责', 'R3, R4, R5, R6', '画像三的"退出和履约风险"'],
        ['主体一致', 'R7', '画像三的"合同/门店/收款三方核验"'],
        ['证据留存', 'R12, R13', '画像三的"证据不足/事后失联"'],
        ['促销甄别', 'R14, R15, R16, R17', '画像二的"规则遗漏"+ 画像三的"团购维权链路"'],
    ],
    [2.5, 3, 8.5],
)

# ═══════════════════════════════════════
# 六、数据模型
# ═══════════════════════════════════════
doc.add_heading('六、数据模型', level=1)

doc.add_heading('6.1 核心实体', level=2)

doc.add_heading('资产（Asset）', level=3)
add_code('id, scene, name, storeName, contractName, payeeName')
add_code('totalPrice, totalTimes, usedTimes, giftTimes')
add_code('validityMonths, weeklyFreq, monthlyBudget')
add_code('unlimited(Boolean), noExpiry(Boolean)')
add_code('refundRule, transferRule, pauseRule')
add_code("groupBuyPlatform(''|meituan|dianping|douyin|other)")
add_code('status(active|paused|expired), createdAt, updatedAt')

doc.add_heading('核销记录（WriteOff）', level=3)
add_code('id, assetId, date, hours, note, remainingAfter, createdAt')

doc.add_heading('资料夹（Folder）+ 凭证文件（File）+ 暂停记录（Pause）', level=3)
add_code('Folder: id, assetId, name, note, createdAt')
add_code('File: id, folderId, name, type, size, materialType, dataUrl(base64), mimeType, uploadedAt')
add_code('Pause: id, assetId, type(pause|transfer), start, end, note, createdAt')

doc.add_heading('6.2 存储设计', level=2)
add_table(
    ['Key', '内容', '说明'],
    [
        ['qf_assets', '资产数组', '核心数据'],
        ['qf_writeoffs', '核销记录数组', '按 assetId 筛选'],
        ['qf_pauses', '暂停记录数组', '按 assetId 筛选'],
        ['qf_folders', '资料夹数组', '每资产自动创建 1 个'],
        ['qf_files', '凭证文件数组（含 base64）', '按 folderId 筛选，注意 5-10MB 上限'],
        ['qf_draft', '套餐录入草稿', '每 10 秒自动保存'],
        ['qf_package_images', '录入页图片临时存储', '与 sessionStorage 分流'],
        ['qf_pin_hash', 'PIN 哈希', "btoa('qf_' + pin).slice(0, 32)"],
        ['qf_unlocked', '解锁标记', "'1' = 已解锁"],
    ],
    [4, 5, 5]
)

# ═══════════════════════════════════════
# 七、技术方案
# ═══════════════════════════════════════
doc.add_heading('七、技术方案', level=1)

doc.add_heading('7.1 技术栈', level=2)
add_table(
    ['层级', '选型'],
    [
        ['框架', 'Vue 3（Composition API / <script setup>）'],
        ['构建', 'Vite 5（端口 3000）'],
        ['路由', 'Vue Router 4（Hash 模式）'],
        ['状态管理', 'Pinia（lock store）'],
        ['样式', 'SCSS（全局变量注入 vite.config.js）'],
        ['图标', '内联 SVG（30+ 图标，24×24 viewBox，CSS 继承颜色）'],
        ['存储', 'localStorage + sessionStorage'],
        ['部署', 'GitHub Pages（push main → Actions → gh-pages）'],
    ],
    [4, 10]
)

doc.add_heading('7.2 关键设计决策', level=2)
add_para('• Hash 路由：兼容 GitHub Pages 无服务端路由')
add_para('• sessionStorage 传参：套餐录入 → 决策卡，图片单独存 localStorage 防超限')
add_para('• 动态 import：11 个页面组件全部路由懒加载')
add_para('• 禁止 router.back()：防止浏览器退出弹窗')
add_para('• 全局 Toast：window.__toast("消息") 带 DOM 兜底')

# ═══════════════════════════════════════
# 八、UI 设计规范
# ═══════════════════════════════════════
doc.add_heading('八、UI 设计规范', level=1)

doc.add_heading('8.1 配色方案', level=2)
add_table(
    ['用途', '色值'],
    [
        ['主色（薄荷绿）', '#48A9A6'],
        ['浅底色', '#B8E6E1'],
        ['按压加深', '#9FD8D2'],
        ['正文色', '#245957'],
        ['辅助文字', '#4A7A77 / #638F8D'],
        ['危险色', '#E8686A'],
        ['页面背景', '#F5FAFA'],
        ['卡片边框', '1.5px solid #48A9A6'],
    ],
    [4, 6]
)

doc.add_heading('8.2 组件尺寸规范', level=2)
add_table(
    ['组件', '规范'],
    [
        ['输入框', '高 44px，圆角 12px'],
        ['按钮', '高 44-50px，圆角 6-8px，按压缩放至 96%'],
        ['卡片', '圆角 12-16px，1.5px 边框'],
        ['弹窗', '宽 343px，圆角 16px'],
        ['进度条', '高 4-6px'],
        ['TabBar', '高 56px，含 safe-area-inset-bottom'],
    ],
    [4, 10]
)

doc.add_heading('8.3 设计原则', level=2)
add_para('• 结论先行：决策卡先展示综合结论，再展开细节')
add_para('• 高风险项红色高亮，通过项默认折叠——"只吵醒该吵的"')
add_para('• 高风险时主按钮变为"我已知晓风险，仍要生成"，返回按钮加粗引导修改——"给台阶下"')
add_para('• 全线 SVG 线条图标，无 emoji')
add_para('• 移动端适配 320-428px，可点击区域 ≥44px')

# ═══════════════════════════════════════
# 九、安全与隐私
# ═══════════════════════════════════════
doc.add_heading('九、安全与隐私', level=1)
add_para('对应画像三的信任设计要求：产品本身也会被当作风险对象，必须透明说明数据处理方式。')
add_table(
    ['措施', '实现'],
    [
        ['全量本地存储', '零服务端传输，所有数据仅在用户设备 localStorage'],
        ['PIN 码锁定', 'btoa 哈希存储（不可逆），锁定后金额/次数全局显示 •••'],
        ['三重防护', '启动锁屏 + 运行时底部锁横幅 + 点横幅直通验证'],
        ['忘记 PIN', '确认弹窗 → 清除哈希重置（数据不丢失）'],
        ['用户可控', '支持随时删除全部数据，支持本地导出'],
    ],
    [4, 10]
)
add_para('⚠ 当前限制：换设备需手动迁移，导出/导入功能 [待开发]。', size=9, color=(0xE8,0x68,0x6A))

# ═══════════════════════════════════════
# 十、产品截图索引
# ═══════════════════════════════════════
doc.add_heading('十、产品截图索引', level=1)
add_para('13 张截图，统一宽 390px，标准屏高度 844px，长截图保留完整高度。右下角叠加 10pt 灰色标题。')
add_table(
    ['编号', '页面', '截图表意'],
    [
        ['01', '首页', '到期提醒 + 双卡片入口 + 预付总额 + 场景标签'],
        ['02', '首页 ⊕ 弹窗', '双录入路径设计'],
        ['03', '资产列表', '多卡管理与不同状态区分'],
        ['04', '套餐录入', '模块一展开 + 实时成本预览'],
        ['05', '套餐录入（长截图）', '模块二三展开 + 场景化文案 + 智能匹配'],
        ['06', '决策卡（长截图）', '结论横幅 + 花费算账 + 问题列表'],
        ['07', '决策卡（长截图）', '行动清单 + 六维度风险详情'],
        ['08', '持仓卡', '无限次充卡价值分析'],
        ['09', '核销录入（长截图）', '到店打卡模式 + 历史记录'],
        ['10', '证据资料夹', '8 类材料清单 + 文件列表 + 一键打包'],
        ['11', '我的', '用户中心 + 三栏统计'],
        ['12', '我的（锁定）', '数据遮蔽 + 锁横幅'],
        ['13', '快速录入', '5 字段极简建卡'],
    ],
    [1.2, 4, 8.8],
)

# ═══════════════════════════════════════
# 十一、已知限制与后续方向
# ═══════════════════════════════════════
doc.add_heading('十一、已知限制与后续方向', level=1)

doc.add_heading('11.1 已知限制', level=2)
add_table(
    ['限制', '影响', '状态'],
    [
        ['localStorage 5-10MB 上限', '大量 base64 图片可能超限', '当前用 SVG 占位图，实际使用需提醒用户'],
        ['无服务端', '无法跨设备同步', '设计如此（隐私优先）'],
        ['浏览器清除数据', '所有记录丢失', '[待补] 导出/导入'],
        ['iOS Safari 隐私模式', 'localStorage 可能受限', '[待确认]'],
    ],
    [4, 5, 5]
)

doc.add_heading('11.2 后续方向（可选建议）', level=2)
add_table(
    ['优先级', '功能', '对应画像需求'],
    [
        ['P0', '数据导出/导入（JSON）', '三画像通用：备份与迁移'],
        ['P1', '到期推送提醒', '画像一：遗忘与闲置干预'],
        ['P1', '消费日历视图', '画像二：统一管理可视化'],
        ['P2', 'AI 合同条款识别（OCR）', '画像三：条款翻译与缺失项检测'],
        ['P2', '商户评分/履约记录社区', '画像三：主体可信度参考（需审慎评估隐私）'],
    ],
    [2, 5, 7],
)

# ═══════════════════════════════════════
# 附录
# ═══════════════════════════════════════
doc.add_heading('附录', level=1)

doc.add_heading('A. 路由表', level=2)
add_table(
    ['路径', '页面', 'Tab'],
    [
        ['/home', '首页', '✅'],
        ['/quick-input', '快速录入', '❌'],
        ['/package-input', '套餐录入', '❌'],
        ['/decision-card', '决策卡', '❌'],
        ['/risk-report', '风险报告', '❌'],
        ['/asset-list', '资产列表', '✅'],
        ['/asset-detail', '持仓卡', '❌'],
        ['/write-off', '核销录入', '❌'],
        ['/evidence-folder', '证据资料夹', '✅'],
        ['/folder-create', '新建资料夹', '❌'],
        ['/mine', '我的', '✅'],
    ],
    [4, 4, 2]
)

doc.add_heading('B. 8 类证据材料清单', level=2)
add_table(
    ['序号', '分类', '材料类型', '关键程度'],
    [
        ['①', '交易依据', '合同协议', '★★★'],
        ['①', '交易依据', '付款截图', '★★★'],
        ['②', '宣传承诺', '活动海报', '★★☆'],
        ['②', '宣传承诺', '销售聊天记录', '★★☆'],
        ['③', '履约记录', '核销打卡记录', '★★☆'],
        ['④', '问题记录', '迁店/停业通知', '★☆☆'],
        ['④', '问题记录', '退费沟通记录', '★☆☆'],
        ['⑤', '用户诉求', '退款转卡协商材料', '★☆☆'],
    ],
    [2, 3, 5, 3]
)

doc.add_heading('C. 研究样本说明', level=2)
add_para('本方案书第二章用户画像基于以下研究：', size=9.5, color=(0x63,0x8F,0x8D))
add_para('• 9 份深度访谈（含 8 份完整文字访谈 + 1 份图片型访谈记录）', size=9.5, color=(0x63,0x8F,0x8D))
add_para('• 87 份问卷，其中 52 人为近两年有效预付消费购买者', size=9.5, color=(0x63,0x8F,0x8D))
add_para('• 购买者中 69.2% 为在校本科生/专科生，65.4% 月可支配 1,000-2,999 元', size=9.5, color=(0x63,0x8F,0x8D))
add_para('• 研究范围覆盖健身、课程、美容、餐饮储值、软件会员等跨场景预付消费', size=9.5, color=(0x63,0x8F,0x8D))

# ── 保存 ──
output_path = r'F:\HuaweiMoveData\Users\HUAWEI\Desktop\青付安 功能截图包\青付安产品方案书.docx'
doc.save(output_path)
print(f'Done: {output_path}')
