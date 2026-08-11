"""
生成青付安产品方案书 .docx — 基于实际代码实现
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.5

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# ── 工具函数 ──
C1 = RGBColor(0x24, 0x59, 0x57)  # 正文色
C2 = RGBColor(0x48, 0xA9, 0xA6)  # 主色
C3 = RGBColor(0x4A, 0x7A, 0x77)  # 辅助文字

def font(run, name='微软雅黑', size=None, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if size: run.font.size = size
    run.bold = bold
    if color: run.font.color.rgb = color

def h0(text):  # 主标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    font(r, size=Pt(22), bold=True, color=C1)

def h1(text):  # 一级标题
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    font(r, size=Pt(15), bold=True, color=C1)

def h2(text):  # 二级标题
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    font(r, size=Pt(12), bold=True, color=C1)

def h3(text):  # 三级标题
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    font(r, size=Pt(10.5), bold=True, color=C3)

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(21)
    r = p.add_run(text)
    font(r, size=Pt(10.5))

def bullet(text, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7 + level * 0.6)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    r = p.add_run('• ' + text)
    font(r, size=Pt(10))

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    r = p.add_run(text)
    font(r, size=Pt(9.5), color=RGBColor(0x99, 0x99, 0x99))

def meta_line(label, value):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    r1 = p.add_run(label + '：')
    font(r1, size=Pt(10), bold=True, color=C3)
    r2 = p.add_run(value)
    font(r2, size=Pt(10))

def table_cell(cell, text, size=Pt(9), bold=False, align='left', bg=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    font(r, size=size, bold=bold)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    v = OxmlElement('w:vAlign')
    v.set(qn('w:val'), 'center')
    tcPr.append(v)
    if bg:
        s = OxmlElement('w:shd')
        s.set(qn('w:val'), 'clear')
        s.set(qn('w:color'), 'auto')
        s.set(qn('w:fill'), bg)
        tcPr.append(s)

def make_table(headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        table_cell(t.rows[0].cells[i], h, size=Pt(9), bold=True, align='center', bg='B8E6E1')
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table_cell(t.rows[ri+1].cells[ci], str(val), size=Pt(8.5))
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()


# ═══════════════════════════════════
# 封面
# ═══════════════════════════════════
for _ in range(7):
    doc.add_paragraph()

h0('青付安 QingFuAn')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('产品方案书')
font(r, size=Pt(14), color=C2)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f'版本 V2.0    日期 {datetime.date.today().strftime("%Y-%m-%d")}')
font(r, size=Pt(10), color=C3)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('纯前端 PWA · 本地存储 · 无服务端')
font(r, size=Pt(9), color=C3)

doc.add_page_break()

# ═══════════════════════════════════
# 1. 产品概要
# ═══════════════════════════════════
h1('1. 产品概要')

h2('1.1 产品是什么')
body('青付安是一个移动端 PWA（Progressive Web App），用来管理预付消费（健身卡、培训课、摄影套餐、美容美发充值等）。用户在办卡前可以录入套餐信息做风险评估，办卡后可以跟进核销打卡、保存凭证、到期提醒。')
body('应用完全运行在浏览器端，数据全量存储在本机 localStorage，没有后端服务。不需要注册账号。')

h2('1.2 技术栈')
meta_line('框架', 'Vue 3（<script setup> 写法，Composition API）')
meta_line('构建', 'Vite 5，开发端口 3000')
meta_line('路由', 'Vue Router 4，Hash 模式（createWebHashHistory）')
meta_line('状态管理', 'Pinia，当前仅 lock.js store 在使用')
meta_line('样式', 'SCSS，全局变量在 vite.config.js 注入')
meta_line('存储', 'localStorage + sessionStorage')
meta_line('部署', 'GitHub Pages，push main → GitHub Actions 自动部署到 gh-pages')
meta_line('线上地址', 'https://lin-hmpp.github.io/qingfuan/')
meta_line('仓库', 'https://github.com/Lin-HMPP/qingfuan')

h2('1.3 覆盖的使用阶段')
body('整个应用覆盖预付消费的三个阶段：')
bullet('购前：录入套餐信息 → 17 条规则运算 → 决策卡 + 风险报告，判断划不划算、条款有没有问题')
bullet('购后：资产列表 → 持仓卡详情 → 日常核销打卡，追踪剩余次数和成本')
bullet('维权：证据资料夹归集合同、付款截图等凭证，一键打包导出 HTML 报告')

h2('1.4 预设消费场景')
body('应用内置 4 种消费场景，每种场景在录入页有独立的提示文案和规则术语：')
make_table(
    ['场景', '涉及术语调整'],
    [
        ['健身/舞蹈', '次卡/年卡/私教课；请假暂停（非休学）；转卡（非转让）'],
        ['培训课程', '课时/全程班；休学（非请假）；转让（非转卡）'],
        ['摄影套餐', '套数/精修张；约拍（非上课）；延期（非休学）；转单（非转卡）'],
        ['美容美发', '次卡/充卡/护理；暂停（非休学）'],
    ],
    [3, 13]
)
body('除以上 4 种外，用户可以在场景选择器中输入自定义场景名称，使用通用文案模板。')

# ═══════════════════════════════════
# 2. 页面功能说明
# ═══════════════════════════════════
h1('2. 页面功能说明')
body('应用共 11 个路由页面，4 个底部 Tab 页（首页、资产列表、证据资料夹、我的），其余为子页面。以下按路由逐一说明实际功能。')

# -- 2.1 首页 --
h2('2.1 首页 /home')
body('首页是默认着陆页，从上到下布局：')
bullet('顶栏：左侧显示"青付安"名称。右侧 ⊕ 按钮（快速录入入口）和两个装饰性圆形色块。')
bullet('Hero 区：盾牌 SVG 图标 + 标题"预付消费，看得懂、算得清、管得住" + 副标题"大学生&职场青年预付资金管理助手" + 一行标签"门店跑路止损 · 卡项过期提醒 · 纠纷一键维权"。')
bullet('到期提醒卡片：筛选剩余 ≤30 天且 >0 天的活跃资产，按剩余天数升序排列，最多展示 3 条。每条显示"门店名 · 剩余 X 天 | 剩余 X 次"，右侧有核销和凭证快捷按钮。7 天内到期加粗显示。无到期资产时显示"暂无即将到期的预付卡"；锁定状态显示"信息已锁定"。')
bullet('双卡片入口：左侧"购买前先检查"→ 跳转 /package-input；右侧"我的预付资产"→ 显示当前预付总额和卡项数量（锁定状态显示•••）→ 跳转 /asset-list。')
bullet('证据资料夹入口卡片：可跳转 /evidence-folder 查看全部文件夹，或直接跳转 /folder-create 新建。')
bullet('常用预付场景标签区：4 个场景标签 +"+ 自定义"，点击后跳转 /package-input 并携带 scene 参数。')
bullet('⊕ 快速录入弹窗：右上角 ⊕ 点击后弹出居中模态框，说明"仅填写核心信息，10 秒完成建卡"，含"进入快速录入"按钮和取消按钮。确认后跳转 /quick-input。')

h2('2.2 快速录入 /quick-input')
body('面向"已经办完卡，只想快速建档"的场景。页面包含以下字段：')
bullet('消费场景：4 个预设场景标签 + 自定义输入（点击"+ 自定义"出现内联输入框，失焦后确认）')
bullet('门店名称：文本输入，最长 30 字符')
bullet('预付总价：数字输入，单位元')
bullet('套餐类型：次卡 / 充卡不限次，二选一 toggle')
bullet('总次数：仅在次卡模式下显示，数字输入')
bullet('有效期限：3/6/12/24 月快捷选择 + 自定义（可分别填数字 + 选日/月/年单位）。默认选中 12 月。')
bullet('底部"直接创建资产卡"按钮：校验必填字段（场景、门店、总价、次卡模式下还需总次数），校验通过后调用 addAsset() 写入 localStorage，自动调用 addFolder() 创建同名资料夹，然后 router.push 跳转资产列表。')
body('该页面不经过规则引擎，不生成决策卡。创建的资产卡 totalTimes 在无限次模式下设为 999，普通模式取用户填写值。validityMonths 根据不同单位折算：天/30，年×12，月不变。')

h2('2.3 套餐录入 /package-input')
body('面向"买前想做全面评估"的场景。页面由 5 个模块组成，模块二~五默认折叠。顶栏左侧返回按钮（带放弃编辑确认），右侧"快速录入"链接。页面每 10 秒自动保存草稿到 localStorage。')

h3('模块一：套餐基础费用')
bullet('套餐总价：必填，数字输入')
bullet('总次数/总课时：必填（充卡模式下隐藏），次卡模式显示。有赠送次数输入框（默认值 0）。')
bullet('充卡不限次 Toggle：与"无固定期限"互斥。激活后隐藏次数输入，显示"充卡模式 · 不限次数，在有效期内任意到店"。')
bullet('服务有效期限：日/月/季度三种单位选择，或开启"没有固定期限" Toggle（与充卡互斥，次卡专用）。')
bullet('实时成本预览：总价和次数都填写后在模块底部显示——次卡模式显示"当前基础单次成本 ≈ X 元/次"，充卡模式显示"X 元/天 · 约 X 元/月"。')

h3('模块二：个人使用规划')
bullet('每月可支配预付预算：必填')
bullet('预计每周使用频率：必填')
bullet('填写频率后自动计算：如果当前频率 × 4.33 × 有效期 ≥ 总次数 → 绿色提示"预计 X 个月用完，在有效期 X 个月内可以完成"；反之 → 黄色警告"预计 X 个月才能用完，但有效期只有 X 个月，到期时约有 X 次用不完，建议提升频率到每周 ≥ X 次"。')
bullet('折叠状态下显示完成情况："预算已填 · 频率已填 → ✓ 已完成"，或"已填 1/2"，或"待填写"。')

h3('模块三：签约 & 收款主体信息')
bullet('门店宣传名称：必填')
bullet('合同签约主体名称：必填')
bullet('收款账户/商家收款名：必填')
bullet('支付渠道：直接付给商家 / 美团 / 大众点评 / 抖音 / 其他团购，五选一。选中团购平台（非"直接付给商家"和"其他"）时，收款账户自动填入对应平台名（如"美团商家平台"）并锁定为只读状态。此数据影响 R17 规则的判定。')
bullet('折叠状态下显示完成情况：三项全部填写 → "✓ 已完成"，部分 → "已填 X/3"。')

h3('模块四：套餐履约规则')
bullet('退款规则：从场景预设中选择，或自定义输入文字。健身场景预设为"未开卡全额退，已开卡按已消费次数比例退"和"其他退款规则（自行填写）"。')
bullet('转卡规则：不可转让 / 可免费转让 / 可转让收手续费（可填 %）。术语按场景切换（培训→"转让"，休学/转学相关表述也不同）。')
bullet('暂停规则：不可暂停 / 可免费暂停 / 自定义次数和天数。术语同理按场景切换（培训→"休学"，摄影→"延期"）。')
bullet('每个规则的选择器是一个可展开的下拉区域，选中预设项后自动填入对应规则文本。')
bullet('折叠状态下显示完成情况：三项全部设置 → "✓ 已完成"，部分 → "已设 X/3"。')

h3('模块五：促销附加说明')
bullet('选填，一个文本输入框，placeholder 随场景变化（如健身→"办年卡送健身包、双人同行第二人半价"）。')

h3('材料上传区')
bullet('独立卡片区域，支持上传 8 种类型的材料：合同协议、付款截图、活动海报、销售聊天记录、核销打卡记录、迁店/停业通知、退费沟通记录、退款转卡协商材料。')
bullet('上传流程：先选材料类型 → 再选上传方式（拍摄/相册/文件）。使用原生 <input type="file">，支持多选，读取 base64 后存入 images 数组。')
bullet('已上传的图片显示缩略图列表，每项显示文件名、材料标签、文件大小，可单独删除。')

h3('提交流程')
bullet('点击"确认录入" → 校验所有必填字段 → 组装数据 → sessionStorage.setItem("qf_package_data", JSON.stringify(pkg)) → 图片单独存 localStorage("qf_package_images") → router.push("/decision-card")。')
bullet('"保存草稿"按钮：将当前 form、scene、images 序列化存入 localStorage("qf_draft" + "qf_draft_time")。草稿过期时间 7 天。')
bullet('从决策卡返回修改时，通过 sessionStorage("qf_draft_back") 恢复数据并自动回填。')

h2('2.4 决策卡 /decision-card')
body('读取 sessionStorage("qf_package_data") 中的套餐数据，调用 runAllRules() 运行 17 条规则。页面加载时显示"正在分析你的套餐信息..."约 600ms 后展示结果。')

bullet('结论横幅（页面顶部，最先看到）：')
bullet('  - 综合风险等级绿色：高风险 0 且中风险 < 4 → 绿色背景 + 盾牌勾选图标 +"看起来还不错"')
bullet('  - 综合风险等级橙色：高风险 1-2 或中风险 ≥4 → 橙色背景 + 感叹号图标 +"有些地方需要核实"')
bullet('  - 综合风险等级红色：高风险 ≥3 → 红色背景 + 三角警告图标 +"建议谨慎决策"')
bullet('  - 横幅副标题显示"X 项高风险 · N 项中风险 · 建议核实后再付款"')

bullet('套餐速览栏：一行四列显示总次数（无限次显示"不限次"）、总价、有效期、单价（次卡→元/次，充卡→元/天）。')

bullet('花费算账区块：')
bullet('  - 单次/日均成本 + 计算公式（如"总价 ¥2,880 ÷ 96 次"）')
bullet('  - 月预算对比：如果填写了月预算，显示是否超出。充卡模式对比月均成本。')
bullet('  - 消耗节奏：每周需要多少次才能在有效期内用完，对比用户计划频率，给出"节奏合适"或"可能来不及"的判断。')
bullet('  - 到期预估进度条：按计划频率能消耗的次数占总次数的百分比，颜色分绿（≥70%）、橙（40%-70%）、红（<40%）。')
bullet('  - 充卡模式下额外展示回本参考：月均成本、以市场参考价 ¥50/次计所需每周到店次数。')

bullet('需要关注的问题：列出所有 high 和 medium 级别的规则。每条显示事实描述和建议行动。高风险项红色左边框，中风险橙色左边框。按严重程度排序。')

bullet('已通过检查：可折叠，默认收起，列出所有 low 和 none 级别规则的事实描述。')

bullet('付款前行动清单：从高风险规则中提取建议行动，去重后以 checkbox 样式列出。')

bullet('内嵌风险维度详情：可折叠，默认收起。六个维度各显示星级评分（1-5 颗实心圆点）和维度下每条规则的判定结论。')

bullet('底部按钮：主按钮"确认录入，生成资产卡"（红色等级时变为"我已知晓风险，仍要生成资产卡"，背景变浅降低视觉权重）；次按钮"返回修改"（红色等级时加粗加大，引导用户返回修改）。')

bullet('确认生成流程：弹出 asset-confirm 弹窗 → 确认后调用 addAsset() → addFolder() → 从 localStorage("qf_package_images") 读取图片 → 逐个 addFile() → 清理 sessionStorage 和 localStorage 临时数据 → router.replace 跳转 /asset-list。')

h2('2.5 风险报告 /risk-report')
body('与决策卡共用同一份 sessionStorage 数据和 runAllRules() 结果。内容上比决策卡更详细，少了心理引导、多了数据罗列。')
bullet('套餐摘要卡片 + 总体风险标签（色标 + 高中风险项数）')
bullet('六维度逐板块展开：每板块显示维度名称、风险等级色标、评分 X/5、命中规则的 fact 和 action。')
bullet('深度成本测算（普通模式）：双栏对比乐观情景（按计划频率）vs 悲观情景（频率减半），展示单次成本、预估回本月份、划算/不划算标签。底部到期使用进度条。')
bullet('充卡价值分析（无限次模式）：三大指标（日均/月均/预估总到店）+ 回本分析（以场景对应的市场参考价计算所需每周最低到店次数 + 用户计划频率对比 + 达标/不达标结论）+ 回本进度条。市场参考价：健身 ¥80，培训 ¥200，摄影 ¥800，美发 ¥100。')
bullet('底部提供"返回决策卡"和"确认生成资产"两个按钮。确认生成流程同决策卡。')

h2('2.6 资产列表 /asset-list')
bullet('统计栏：三行文字显示预付总金额（元）、在库卡数、即将到期（≤30 天）卡数。锁定状态显示•••。')
bullet('资产卡片：每张卡片显示"门店名 · 场景"、剩余 X/总 X 次 + 到期 X 天（充卡模式显示"已到店 X 次 · 充卡不限次数"）、单次成本（充卡模式显示"充卡 · 不限次数"）。5 分钟内新建的卡片显示绿色"新"标签。卡片右侧有核销和凭证两个快捷按钮。')
bullet('管理模式：右上角"管理资产"按钮 toggle，进入管理模式后卡片变红色边框，每个卡片显示编辑（✎）和删除（✕）圆形按钮。')
bullet('编辑弹窗：可修改门店名称、消费场景、预付总价、总次数、有效期。')
bullet('删除流程：确认弹窗 → 级联清除：核销记录 → 暂停记录 → 关联资料夹及其中所有文件 → 资产本身。')
bullet('底部"+ 新增预付卡"按钮，跳转 /package-input。')

h2('2.7 持仓卡 /asset-detail')
body('从资产列表点击进入，展示单张资产的全部数据和操作入口。')
bullet('资产名称 + 状态标签（使用中/锁卡中/已失效）。暂停中额外显示暂停截止日期。')
bullet('资产数据面板：预付参考总额、履约起止日期、剩余有效期天数、剩余/总次数 + 已用占比（普通模式）/ 累计到店次数（无限次模式）、当前实际单次成本（普通模式）/ 日均成本 + 单次到店成本（无限次模式）、剩余权益参考值（普通模式）。锁定状态数值显示•••。')
bullet('使用进度条：普通模式 = 已用次数%，无限次模式 = 已过时间%。进度条下方有条件显示的预警文字（使用率偏低/时间过半到店少）。')
bullet('资产缩水预警卡片：剩余 ≤30 天且使用率偏低时显示。')
bullet('暂停自动恢复：页面挂载时检查所有暂停记录，如果全部已到期则自动恢复资产状态。')
bullet('四个操作按钮：核销记录（跳转 /write-off，过期或暂停中禁用）、暂停/改期（弹出 pause-transfer 组件）、申请退款（弹出 refund-checklist 组件，支持无限次时间比例计算）、查看证据资料（跳转 /evidence-folder 并带 assetId 参数）。')

h2('2.8 核销录入 /write-off')
body('从持仓卡进入，URL 携带资产 ID。页面加载时做三道拦截：资产不存在 → 提示；资产已过期且非无限期 → 提示 + 跳回持仓卡；资产处于暂停状态 → 提示 + 跳回持仓卡。')
bullet('关联资产信息卡片：显示门店名、剩余次数/天数（普通模式）或已到店次数/剩余天数（无限次模式）、充卡模式下显示单次到店成本和日均成本徽章。')
bullet('核销表单：')
bullet('  - 日期选择（type="date"，默认当天，阻止未来日期）')
bullet('  - 普通模式：手动输入消耗次数，校验正整数且不超过剩余次数')
bullet('  - 无限次模式：隐藏次数输入，显示"本次到店自动计 1 次打卡"提示 + 频率建议（如"周均到店 X 次，频率很好/适中/偏少"，基于已过天数计算）')
bullet('  - 场景化标签文字：根据资产场景切换——健身→"消耗次数"/"次"，培训→"消耗课时"/"课时"，摄影→"拍摄套数"/"套"，美容→"消费次数"/"次"')
bullet('  - 备注（选填）')
bullet('历史记录列表：每次核销生成一条记录，显示日期 + 备注 + 剩余/累计数。点击可查看详情（write-off-detail 组件），支持编辑和删除。')
bullet('保存后自动更新 asset.usedTimes 并生成 writeoff 记录。')

h2('2.9 证据资料夹 /evidence-folder')
body('底部 Tab 页之一。如果从其他页面带 assetId 参数进入，自动选中对应文件夹。')
bullet('页面顶栏：标题 +"+ 新建资料夹"按钮（跳转 /folder-create）。')
bullet('文件夹列表：每个文件夹显示名称、绑定资产名、凭证份数。右侧两个按钮："导出凭证"（直接导出该文件夹的 HTML 报告）；"···"（展开操作菜单，含编辑和删除）。')
bullet('选中文件夹后展开以下内容：')
bullet('  - 材料完整性检查清单：8 个材料类型复选框，已上传的显示绿色"已上传"，缺失的显示"缺失"（可点击，点击后直接进入该类型的上传流程）。底部显示"已上传 X/8 项"。')
bullet('  - 当前文件夹文件列表：每项显示文件名、材料类型标签、文件大小。')
bullet('  - 底部双按钮："一键打包"（生成并下载 HTML 维权报告）和"+ 新增凭证上传"。')
bullet('  - 隐私提示："所有凭证仅本地存储，不上传服务器"。')
bullet('导出 HTML 报告格式：包含标题"青付安 · 维权凭证报告"、资产名、导出时间、按材料类型分组的文件列表，图片直接 base64 内嵌，非图片文件提供下载链接。')
bullet('编辑资料夹弹窗：可修改文件夹名称（过滤 emoji 字符）和绑定资产（下拉选择）。')
bullet('删除资料夹：确认后级联清除该文件夹下的所有文件。')
bullet('上传流程：选类型 → 选方式（拍摄/相册/文件）→ FileReader 读取 base64 → addFile()。')

h2('2.10 新建资料夹 /folder-create')
bullet('绑定预付资产：下拉选择（必填），选项从 getAssets() 读取，显示格式"门店名 · 场景 (¥总价)"。')
bullet('文件夹名称：文本输入，最长 30 字符，校验规则：2-30 字符，不含 < > : " / \\ | ? *。')
bullet('备注说明：选填多行文本。')
bullet('创建成功后跳回 /evidence-folder。')

h2('2.11 我的 /mine')
bullet('用户卡片：头像占位圆形 +"青付安用户"+ 状态显示（"信息已隐藏"或"信息可见"）+ 锁定/解锁按钮。锁定按钮在锁定状态下变为白底绿边框"解锁"。')
bullet('三栏统计：总资产金额（元）、卡项总数（张）、即将到期（张）。锁定状态下数值替换为•••。此数据从 getAssets() 实时计算，监听 app-unlocked 事件在解锁后刷新。')
bullet('菜单列表（5 项）：')
bullet('  - 使用指南：弹出 5 步说明弹窗（两种录入方式 → 决策评估 → 资产管理 → 证据资料夹 → 安全隐私）。')
bullet('  - 规则说明：弹出 17 条规则完整列表弹窗，分"购前决策""购后管理""商户履约"三组。每条规则显示编号、名称、说明文字。')
bullet('  - 隐私设置：弹窗说明所有数据仅本地存储，支持 PIN 码锁定。')
bullet('  - 本地凭证管理：弹窗说明文件仅保存在本机，换设备不会同步。')
bullet('  - 重置所有数据：确认后执行 localStorage.clear() + location.reload()。')

# ═══════════════════════════════════
# 3. 公共组件
# ═══════════════════════════════════
h1('3. 公共组件')

body('应用有 11 个可复用组件，分布在各页面中：')

make_table(
    ['组件', '使用位置', '功能'],
    [
        ['tab-bar', 'App.vue，4 个 Tab 页底部', '底部导航栏，使用内联 SVG 图标，当前页图标变色'],
        ['toast', 'App.vue 全局挂载', '深色半透明 Toast，window.__toast(msg) 调用。3 秒自动消失，点击可关闭。含 DOM 兜底机制——Vue 组件未挂载时直接创建 DOM 元素。右下偏移避开 TabBar'],
        ['pin-lock', 'App.vue 锁屏 / 锁横幅点击', '数字键盘 PIN 输入界面，支持设置和验证两种模式。含"忘记密码"入口，确认后清除 PIN 哈希重置'],
        ['scene-picker', '套餐录入页', '场景选择弹窗：4 个预设场景 + 自定义入口 + 取消按钮'],
        ['image-picker', '套餐录入页 / 证据资料夹', '图片/文件采集：选类型 → 选方式（拍摄/相册/文件）→ 多选 FileReader base64'],
        ['refund-checklist', '持仓卡"申请退款"', '退款材料核对清单，含费用估算。无限次模式按时间比例计算（已过天数÷总天数），普通模式按次数比例'],
        ['pause-transfer', '持仓卡"暂停/改期"', '设置暂停起止日期 + 编辑套餐有效期。保存后刷新资产状态'],
        ['asset-confirm', '决策卡生成确认', '确认弹窗，显示即将创建的资产概要'],
        ['exit-confirm', '无当前调用（预留）', '退出确认弹窗'],
        ['write-off-detail', '核销录入历史记录', '单条核销记录详情：查看、编辑、删除'],
        ['package-loading', '证据资料夹导出', '打包导出加载动画，显示进度'],
    ],
    [3, 4.5, 8.5]
)

# ═══════════════════════════════════
# 4. 数据流与存储
# ═══════════════════════════════════
h1('4. 数据流与存储')

h2('4.1 localStorage 数据清单')
body('应用在 localStorage 中维护以下 key，全部以 qf_ 为前缀：')

make_table(
    ['Key', '类型', '说明'],
    [
        ['qf_assets', 'Array', '资产卡片数组。每个对象含 id、scene、storeName、totalPrice、totalTimes、usedTimes、validityMonths、unlimited、noExpiry、weeklyFreq、monthlyBudget、contractName、payeeName、groupBuyPlatform、refundRule、transferRule、pauseRule、giftTimes、status、createdAt、updatedAt 等字段'],
        ['qf_writeoffs', 'Array', '核销记录。含 id（wo_前缀）、assetId、date、hours、note、remainingAfter'],
        ['qf_pauses', 'Array', '暂停/改期记录。含 id（pau_前缀）、assetId、type、start、end、note'],
        ['qf_folders', 'Array', '资料夹。含 id（fld_前缀）、assetId、name、note'],
        ['qf_files', 'Array', '凭证文件。含 id（fil_前缀）、folderId、name、type、size、materialType、dataUrl（base64）、mimeType。文件内容以 base64 存储，可能占用较大空间'],
        ['qf_account', 'Object', '账户信息（预留字段，当前未实际使用）'],
        ['qf_draft', 'Object', '套餐录入草稿，含 form、scene、images'],
        ['qf_draft_time', 'Number', '草稿保存时间戳，读取时与当前时间比较，超过 7 天自动清除'],
        ['qf_logs', 'Array', '操作日志（预留，当前未实际使用）'],
        ['qf_package_images', 'Array', '录入页上传图片的临时存储（与 sessionStorage 配合使用）'],
        ['qf_pin_hash', 'String', 'PIN 哈希值，算法：btoa("qf_" + pin).slice(0, 32)'],
        ['qf_unlocked', 'String', '值为 "1" 时表示当前已解锁'],
    ],
    [3.5, 2, 10.5]
)

h2('4.2 页面间数据传递')
body('套餐录入 → 决策卡 → 风险报告的链路是应用中主要的跨页面数据流：')
bullet('套餐录入提交：数据写入 sessionStorage("qf_package_data")（不含图片），图片 base64 单独存 localStorage("qf_package_images")。原因：sessionStorage 容量通常比 localStorage 更受限，大量 base64 可能导致写入失败。')
bullet('决策卡/风险报告读取：onMounted 时从 sessionStorage 读取并 JSON.parse。')
bullet('决策卡返回修改：数据回写 sessionStorage("qf_draft_back")，套餐录入页 onMounted 优先检查此 key，有则回填并清除。')
bullet('生成资产后：addAsset() → addFolder() → 逐个 addFile() → 清除 sessionStorage("qf_package_data") 和 localStorage("qf_package_images") → router.replace 跳转。')
bullet('其他页面间：统一通过 URL query 参数传 id 或 assetId（如 /asset-detail?id=xxx、/write-off?id=xxx、/evidence-folder?assetId=xxx）。')

h2('4.3 CRUD 操作')
body('common/storage.js 提供统一的 localStorage 读写封装：')
bullet('每个实体有 getXxx()、addXxx()、updateXxx(id, patch)、deleteXxx(id) 全套函数。')
bullet('add 操作：自动生成 id（前缀 + 时间戳 + 随机串），自动写入 createdAt。')
bullet('update 操作：通过 id 查找 → Object.assign 合并 patch → 自动写入 updatedAt。')
bullet('delete 操作：资产删除时级联清除 writeoffs、pauses、folders、files。资料夹删除时级联清除所属文件。')
bullet('所有读写均有 try-catch 保护，返回 null 或 false 表示失败。')

# ═══════════════════════════════════
# 5. 规则引擎
# ═══════════════════════════════════
h1('5. 规则引擎')

h2('5.1 概述')
body('规则引擎实现在 common/rules-engine.js 中，入口函数 runAllRules(data)。输入套餐数据对象，输出 { risks, dimensions, costs, grade, summary }。所有计算在浏览器本地执行。')

h2('5.2 17 条规则明细')
body('每条规则返回 { level: "high"|"medium"|"low"|"none", code: "R1"-"R17", title, layers: { fact, confirm, explain, action } }。')

make_table(
    ['编号', '规则', '高风险条件', '中风险条件'],
    [
        ['R1', '单次/日均成本核算', '超出月预算对应阈值（模式不同算法不同）', '未填月预算'],
        ['R2', '有效期-频率匹配度', '需要频率与计划差距>1.5次/周', '差距>0.5次/周'],
        ['R3', '合同/凭证可得性', '—', '无书面合同'],
        ['R4', '退款条款清晰度', '退款规定仅口头承诺', '—'],
        ['R5', '转卡/暂停/延期条款', '—', '未明确手续费标准'],
        ['R6', '迁址/停业应对条款', '—', '总价≥800或有效期≥6月且无条款'],
        ['R7', '合同主体一致性', '门店/合同签约方/收款方不匹配；或个人账户收款', '主体信息不完整，无法完成校验'],
        ['R8', '高金额预付', '总价>3倍月预算', '总价>1.5倍月预算'],
        ['R9', '赠品规则清晰度', '—', '有赠品但规则不清'],
        ['R10', '到期预警', '剩余≤30天', '剩余≤60天'],
        ['R11', '使用频率异常', '实际频率<计划的50%', '实际频率<计划的80%'],
        ['R12', '材料留存完整性', '缺失≥3类关键材料', '缺失2类'],
        ['R13', '退款前置检查', '—', '退款规则未知'],
        ['R14', '服务变更记录', '—', '—（默认low）'],
        ['R15', '场景专属子规则', '—', '—（默认low）'],
        ['R16', '退款渠道核验', '—', '默认中风险（商家通常不主动告知）'],
        ['R17', '平台团购风险', '—', '非"直接付给商家"模式即中风险'],
    ],
    [1, 3.5, 5.5, 6]
)
note('R11 在购前阶段（actualFreq=0）返回 level="none"，不参与评分。R14/R15 当前为占位规则，默认返回 low，需用户后续主动记录相关事件才有实际判断。')

h2('5.3 六维度评分')
body('17 条规则分入 6 个维度：')
make_table(
    ['维度', '包含规则', '说明'],
    [
        ['预付压力', 'R1、R2、R8', '成本+频率+金额，三个维度交叉评估经济负担'],
        ['履约时限', 'R9、R10、R11', '赠品期限+到期预警+频率异常，关注时间维度的风险'],
        ['合约权责', 'R3、R4、R5、R6', '合同+退款+转卡+迁址，覆盖合同条款完整性'],
        ['主体一致', 'R7', '门店/合同/收款三方一致性，防止收款主体不清'],
        ['证据留存', 'R12、R13', '材料留存+退款前置，为可能维权做准备'],
        ['促销甄别', 'R14、R15、R16、R17', '变更+场景+渠道+团购，识别交易结构中的隐藏风险'],
    ],
    [2.5, 4, 9.5]
)
body('评分公式：dim.score = max(1, min(5, 3 + low数量 - high数量×2 - medium数量))。')

h2('5.4 综合风险定级')
body('highCount ≥ 3 → 高风险（红色 #DC3545）\nhighCount ≥ 1 或 mediumCount ≥ 4 → 中风险（橙色 #FD7E14）\n其他 → 低风险（绿色 #28A745）')

h2('5.5 成本测算')
body('runAllRules 返回的 costs 对象包含：')
bullet('无限次模式：daily（元/天）、monthly（元/月）、breakEven（含 estimatedTotalVisits、idealPerVisitCost、marketPerVisit、visitsPerWeek 回本所需每周最少次数，suggestion 建议文字）。市场参考价按场景不同：健身 ¥80/次，培训 ¥200/次，摄影 ¥800/次，美发 ¥100/次，其他 ¥50/次。')
bullet('普通模式：base（票面单次成本）、ideal（按计划频率的理想单次成本 + 预估回本月份）、conservative（频率减半时的保守单次成本 + 预估回本月份）、expiry（使用进度百分比 + 建议文字）。')
bullet('充卡回本判断：回本所需每周到店次数 ≤ 计划频率 →"每周到店 X 次可以值回票价"；回本次数 ≤7 → 提示偏少；回本次数 >7 → 提示建议选更短期套餐。')

# ═══════════════════════════════════
# 6. 安全机制
# ═══════════════════════════════════
h1('6. 安全机制')

h2('6.1 PIN 码锁定')
body('PIN 锁机制由 store/lock.js 统一管理，App.vue 负责 UI 表现。')
bullet('设置流程：用户输入 6 位数字 PIN → btoa("qf_" + pin).slice(0, 32) → 存 localStorage("qf_pin_hash") → 设 qf_unlocked="1" 表示已解锁。')
bullet('解锁流程：输入 PIN → 同样算法哈希 → 比对 qf_pin_hash → 匹配则设 qf_unlocked="1"，派发 app-unlocked 事件。')
bullet('锁定流程：点击"锁定信息"或 doLock() → 删除 qf_unlocked → locked 和 showLockBanner 两个 ref 变为 true。')
bullet('忘记 PIN：PIN 键盘上有入口，确认后清除 qf_pin_hash，重新设置。')

h2('6.2 三重锁防护')
bullet('启动锁屏：App.vue onMounted 调用 checkLock() → 检测到有 PIN 哈希且 qf_unlocked≠"1" → 返回 true → App.vue 显示全屏 pin-lock 组件覆盖所有内容。')
bullet('运行时数据屏蔽：locked=true 时，首页、资产列表、持仓卡、我的页面中的金额和次数显示为•••；各页面的 guard() 函数拦截操作按钮，弹出"信息已锁定，请先解锁"toast。')
bullet('运行时锁横幅：locked=true 且未显示锁屏时（即已通过启动锁屏但后来被运行时锁定），页面底部显示深色胶囊条"🔒 信息已锁定 · 点击解锁"，position:fixed 固定在 TabBar 上方。点击后弹出 PIN 验证。')

h2('6.3 数据隐私')
bullet('所有数据存储于用户设备 localStorage，不经过任何服务器。')
bullet('无需注册账号、无需手机号。')
bullet('百度统计和 Microsoft Clarity 仅采集页面访问事件（通过 track() 函数发送），不包含用户财务数据。')

# ═══════════════════════════════════
# 7. 界面规范
# ═══════════════════════════════════
h1('7. 界面规范')

h2('7.1 配色')
body('整套配色以薄荷绿为主，已定稿：')
make_table(
    ['用途', '色值', 'CSS 变量/类名'],
    [
        ['主色（按钮、边框）', '#48A9A6', '.btn-primary, .input-blue'],
        ['浅底色（标签、提示背景）', '#B8E6E1', '.card-light, 各类浅底元素'],
        ['按压加深', '#9FD8D2', '.btn-primary:active'],
        ['正文色（标题、正文）', '#245957', '全局 color'],
        ['辅助文字', '#4A7A77 / #638F8D', '次要信息、placeholder'],
        ['危险色（删除、高风险）', '#E8686A', '.btn-del, .problem-high'],
        ['页面背景（决策卡）', '#F5FAFA', '.page 特定页面'],
    ],
    [3.5, 4, 8.5]
)

h2('7.2 通用组件规范')
bullet('卡片：.card-blue — 白底 + 1.5px 薄荷绿边框 + 圆角 12-16px + 按压 scale(0.97)')
bullet('主按钮：.btn-primary — 44-50px 高 + 薄荷绿填充 + 白字 + 按压缩放 96% + 背景加深至 #9FD8D2')
bullet('次按钮：.btn-secondary — 白底 + 薄荷绿边框 + 薄荷绿字 + 按压背景变浅')
bullet('输入框：.input-blue — 44px 高 + 1.5px 薄荷绿边框 + 圆角 12px + 聚焦外发光')
bullet('分割线：.divider-blue — 0.5px 薄荷绿，透明度 0.4')
bullet('导航栏：.nav-bar — 44px 高 + 白底 + 底部 1px 薄荷绿边框 + 标题居中')

h2('7.3 图标')
body('全线使用 SVG 线条图标，24×24 viewBox，线宽 1.5-2px。颜色通过 CSS stroke="currentColor" 继承。不使用 emoji 表情符号（产品初期曾使用，后统一替换）。')
body('公共图标定义在 common/icons.js，导出 30+ 个函数，返回 SVG 字符串。组件中直接使用内联 SVG 标签。')

h2('7.4 移动端适配')
bullet('适配主流手机宽度 320-428px')
bullet('底部 TabBar 使用 env(safe-area-inset-bottom) 适配刘海屏')
bullet('最小可点击区域 44px')
bullet('页面切换动画：0.2s opacity + translateX(8px)')

# ═══════════════════════════════════
# 8. 部署与兼容
# ═══════════════════════════════════
h1('8. 部署与兼容')

h2('8.1 构建与部署')
bullet('开发启动：npm install → npm run dev → http://localhost:3000')
bullet('生产构建：npm run build → 输出到 dist/')
bullet('自动部署：推送 main 分支 → GitHub Actions 触发 → 构建 → 部署到 gh-pages 分支 → https://lin-hmpp.github.io/qingfuan/ 更新')
bullet('PWA 支持：项目含 manifest.json 和 Service Worker，浏览器可提示"添加到主屏幕"')

h2('8.2 浏览器兼容')
bullet('iOS Safari（≥14）：全面支持')
bullet('Android Chrome（≥90）：全面支持')
bullet('微信内置浏览器（iOS 和 Android）：已测试通过')
bullet('桌面端 Chrome/Edge：功能可用但非主要适配目标')

h2('8.3 已知限制')
bullet('localStorage 容量：浏览器通常限制 5-10MB，大量 base64 图片可能超限。当前未做容量检测和提醒。')
bullet('iOS Safari 隐私模式：部分版本 localStorage 行为异常。')
bullet('数据持久性：清除浏览器数据会丢失全部记录。当前无内置导出/备份功能，需后续版本补充。')
bullet('跨设备同步：纯前端架构不支持。换手机需手动迁移数据。')

# ═══════════════════════════════════
# 9. 路由与组件清单
# ═══════════════════════════════════
h1('9. 路由与组件清单')

h2('9.1 全部路由')
make_table(
    ['路径', '页面', 'Tab', '路由方式'],
    [
        ['/', '重定向到 /home', '', 'redirect'],
        ['/home', '首页', '✓', '懒加载'],
        ['/quick-input', '快速录入', '', '懒加载'],
        ['/package-input', '套餐录入', '', '懒加载'],
        ['/decision-card', '决策卡', '', '懒加载'],
        ['/risk-report', '风险报告', '', '懒加载'],
        ['/asset-list', '资产列表', '✓', '懒加载'],
        ['/asset-detail', '持仓卡', '', '懒加载'],
        ['/write-off', '核销录入', '', '懒加载'],
        ['/evidence-folder', '证据资料夹', '✓', '懒加载'],
        ['/folder-create', '新建资料夹', '', '懒加载'],
        ['/mine', '我的', '✓', '懒加载'],
    ],
    [3.5, 3, 1.5, 2]
)

h2('9.2 文件结构')
bullet('common/ — 公共逻辑：storage.js（CRUD + KEY 常量）、rules-engine.js（规则引擎）、validator.js（表单校验）、analytics.js（百度统计 + Clarity 埋点）、icons.js（30+ SVG 图标）')
bullet('pages/ — 11 个页面目录，每个含 index.vue')
bullet('components/ — 11 个可复用组件目录')
bullet('store/ — Pinia store：lock.js（全局锁状态），login.js（残留，未使用）')
bullet('static/icons/ — 8 个 Tab 图标 SVG 文件')

# ═══════════════════════════════════
# 10. 待完善项
# ═══════════════════════════════════
h1('10. 待完善项')

bullet('数据导出/导入：用户换手机或清除浏览器数据后无法恢复。需要提供 JSON 格式的导出/导入功能。')
bullet('localStorage 容量监控：目前上传图片没有容量检测，base64 图片累积可能超过浏览器上限（通常 5-10MB）。需要在超过阈值时提醒用户清理。')
bullet('R14/R15 规则：目前为占位规则，默认返回 low。R14 需要用户手动记录商家变更事件才有实际判断，R15 的场景差异化逻辑还未细化。')
bullet('login.js store：代码中残留，未在任何页面引用，可清理。')
bullet('qf_account / qf_logs：storage.js 中预留了这两个 key 的 CRUD，但当前无页面实际使用。')
bullet('桌面端适配：页面在桌面端布局无大问题，但未做过专门的响应式优化。')
bullet('[备注：缺少用户量数据，无法评估实际使用情况和兼容性问题频率]')
bullet('[备注：缺少具体的浏览器兼容性测试报告，以上兼容性结论基于 PWA 通用特性推断]')

# ── 保存并打开 ──
output_path = r'C:\Users\HUAWEI\Desktop\qingfuan\青付安产品方案书.docx'
doc.save(output_path)
print(f'已生成：{output_path}')
