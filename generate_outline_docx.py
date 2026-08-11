"""
生成青付安演示视频录制大纲 Word 文档
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = '微软雅黑'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    hs.font.color.rgb = RGBColor(0x24, 0x59, 0x57)
    hs.font.bold = True
    if i == 1:
        hs.font.size = Pt(22)
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(12)
    elif i == 2:
        hs.font.size = Pt(16)
        hs.paragraph_format.space_before = Pt(20)
        hs.paragraph_format.space_after = Pt(8)
    else:
        hs.font.size = Pt(13)
        hs.paragraph_format.space_before = Pt(14)
        hs.paragraph_format.space_after = Pt(6)

def add_para(text, bold=False, size=10.5, color=None, align=None, space_after=6, indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def set_cell_shading(cell, color_hex):
    tcPr = cell._element.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:shd')):
        tcPr.remove(existing)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, size=9, color=None, align='left'):
    for p in cell.paragraphs:
        p.clear()
    p = cell.paragraphs[0]
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, '48A9A6')
        set_cell_text(cell, h, bold=True, size=9, color=(255,255,255), align='center')
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            if r % 2 == 1:
                set_cell_shading(cell, 'F5FAFA')
            set_cell_text(cell, val, size=9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

# ═══════════════════════ 封面 ═══════════════════════
doc.add_paragraph()
doc.add_paragraph()
add_para('青付安（QingFuAn）', bold=True, size=28, color=(0x24,0x59,0x57), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('演示视频录制大纲', bold=True, size=20, color=(0x48,0xA9,0xA6), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
add_para('总时长：~3 分钟 | 核心策略：痛点共鸣 → 核心能力展示 → 差异化收尾', size=11, color=(0x63,0x8F,0x8D), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
add_para('公开链接：https://lin-hmpp.github.io/qingfuan/', size=10, color=(0x4A,0x7A,0x77), align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ═══════════════════════ 目录 ═══════════════════════
doc.add_heading('目录', level=1)
for item in ['前期准备', '第一幕：痛点钩子（0:00 - 0:20）', '第二幕：核心能力——从录卡到决策（0:20 - 1:50）', '第三幕：差异化——无限次充卡模式（1:50 - 2:20）', '第四幕：收尾——证据夹 + 隐私锁（2:20 - 2:50）', '结尾（2:50 - 3:00）', '设计说明 + 速查表']:
    add_para(item, size=11, space_after=4)

doc.add_page_break()

# ═══════════════════════ 前期准备 ═══════════════════════
doc.add_heading('前期准备', level=1)

add_para('1. 设备设置', bold=True, size=11)
add_para('• Chrome 打开 https://lin-hmpp.github.io/qingfuan/', indent=0.5)
add_para('• F12 → Ctrl+Shift+M → 选 iPhone 14（390×844）', indent=0.5)
add_para('• 控制台粘贴 demo-data.js → 回车 → 刷新页面', indent=0.5)
add_para('• 确认首页显示：到期提醒 + 4 张资产卡 + 场景标签', indent=0.5)

add_para('2. 效率技巧', bold=True, size=11)
add_para('• 在手机备忘录提前打好所有要输入的文字，录屏时复制粘贴，避免键盘遮挡画面', indent=0.5)
add_para('• 输入环节可适当加速（后期 1.2×），实际录制时不必等字打完再说旁白', indent=0.5)
add_para('• 旁白控制在 180-200 字/分钟，全程约 600-700 字', indent=0.5)

add_para('3. 输入内容速查', bold=True, size=11)
add_table(
    ['场景', '字段', '输入值'],
    [
        ['套餐录入 - 模块一', '总价 / 次数 / 赠送 / 有效期', '2880 / 96 / 2 / 12 个月'],
        ['套餐录入 - 模块二', '每月预算 / 每周频率', '500 / 3'],
        ['套餐录入 - 模块三', '门店 / 合同 / 收款', 'XX健身工作室 / XX体育文化发展有限公司 / XX体育文化发展有限公司'],
        ['套餐录入 - 模块三', '支付渠道', '直接付给商家'],
        ['套餐录入 - 模块四', '退款 / 转卡 / 暂停', '各选一个预设项（点选即可）'],
        ['PIN 锁定', 'PIN 码', '111111'],
    ],
    [3, 5, 6]
)

doc.add_page_break()

# ═══════════════════════ 第一幕 ═══════════════════════
doc.add_heading('第一幕：痛点钩子（0:00 - 0:20）', level=1)

add_para('目标：5 秒内让观众产生"这就是我"的共鸣', bold=True, size=10, color=(0xE8,0x68,0x6A))

add_table(
    ['时间', '画面操作', '旁白（建议）'],
    [
        ['0:00', '打开首页，手指指向到期提醒第一条\n"XX健身工作室·套餐 剩余74次 30天后到期"', '"这是很多人的故事——花 2880 办了张健身年卡，去过 22 次，还剩 74 次，还有 30 天到期。"'],
        ['0:10', '快速下划扫过首页全貌，停在预付总额数字上', '"不是不想去，是根本没算过——这一笔亏了多少、下一笔该不该办。"'],
    ],
    [1.5, 5, 7.5]
)

doc.add_page_break()

# ═══════════════════════ 第二幕 ═══════════════════════
doc.add_heading('第二幕：核心能力——从录卡到决策（0:20 - 1:50）', level=1)

add_para('目标：展示最核心的价值链路——信息录入 → 规则分析 → 决策卡结论', bold=True, size=10, color=(0xE8,0x68,0x6A))
add_para('这是全片最重要的 90 秒，决定了评委对产品核心能力的判断。', size=10, color=(0x63,0x8F,0x8D))

add_table(
    ['时间', '画面操作', '旁白（建议）'],
    [
        ['0:20', '点击「购买前先检查」，场景已是"健身/舞蹈"', '"假设现在又有一个办卡冲动——先花一分钟录入信息。"'],
        ['0:25', '快速连续填写：总价 2880 → 次数 96 → 赠送 2 → 12 个月', ''],
        ['0:35', '手指停在实时成本「当前基础单次成本 ≈ 30.0 元/次」', '"总价和次数一填，实时告诉你每次实际成本——别被月均低价骗了。"'],
        ['0:42', '展开模块二，填 500 / 3，手指指向绿色匹配提示', '"你说每周去 3 次，系统算出来：按这个频率 7 个月用完，12 个月有效期内没问题。"'],
        ['0:50', '展开模块三，填门店/合同/收款，重点指支付渠道标签「直接付给商家」', '"但关键来了——签合同的公司、收款方是不是同一家？不一致就是高风险——钱转给个人，出事找谁？"'],
        ['1:00', '展开模块四，快速点选退款/转卡/暂停各一项预设', '"退款、转卡、暂停的规则，点选就行。"'],
        ['1:08', '点击底部「确认录入，生成预付资产卡片」，等待加载', ''],
        ['1:15', '决策卡出现——这是全片最核心画面，停留 3 秒\n缓慢向下滑动：结论横幅 → 套餐速览 → 花费算账', '"17 条规则分析后，出决策卡。结论直接放最上面：红色警告，说清楚问题在哪。下面是花费算账——每次成本、月预算对比、消耗节奏；需要关注的问题，高风险标红。"'],
        ['1:35', '继续下滑到行动清单，展开一个风险维度的详情', '"每个问题都附带行动建议——付款前要核对什么、要商家补充什么，照着做就行。"'],
    ],
    [1.5, 5.5, 7]
)

doc.add_page_break()

# ═══════════════════════ 第三幕 ═══════════════════════
doc.add_heading('第三幕：差异化——无限次充卡模式（1:50 - 2:20）', level=1)

add_para('目标：展示与传统记账工具的差异——充卡/年卡模式的专门适配', bold=True, size=10, color=(0xE8,0x68,0x6A))

add_table(
    ['时间', '画面操作', '旁白（建议）'],
    [
        ['1:50', '点击底部 Tab「资产」→ 点击「XX美发沙龙·套餐」（充卡标签）', '"办了卡之后呢？看这张美发充年卡——一次付 3800，不限次数。"'],
        ['1:55', '依次手指指向：累计到店 16 次 → 日均成本 → 时间进度条', '"不限次数不是不算账——累计到店几次、每天划多少钱、有效期过了多少，全都可视化。"'],
        ['2:05', '点击「核销」→ 手指指打卡界面（自动显示"到店打卡"徽章）', '"到店就点一下，自动打卡。传统记账软件管不了充卡模式，我们专门做了适配。"'],
    ],
    [1.5, 5.5, 7]
)

# ═══════════════════════ 第四幕 ═══════════════════════
doc.add_heading('第四幕：收尾——证据夹 + 隐私锁（2:20 - 2:50）', level=1)

add_para('目标：展示闭环价值——出事时有证据、数据始终在自己手里', bold=True, size=10, color=(0xE8,0x68,0x6A))

add_table(
    ['时间', '画面操作', '旁白（建议）'],
    [
        ['2:20', '点击底部 Tab「证据夹」→ 点击「XX健身工作室 凭证」', '"万一出纠纷——比如门店跑路——每个资产自动建一个证据资料夹。"'],
        ['2:25', '手指指材料检查清单（合同✓ 付款✓ 海报缺失等），点一个"缺失"项跳转上传界面 → 取消返回', '"8 类维权材料清单，有什么缺什么一目了然。缺的直接点，拍照上传。"'],
        ['2:35', '手指停在底部「一键打包」按钮 → 点击（可不真导出）', '"一键打包成报告，合同、付款截图全在里面，可以直接打印提交 12315。出事了才找材料已经晚了，青付安让你办卡那一刻起就开始留证据。"'],
        ['2:40', '点击底部 Tab「我的」→ 点击「锁定信息」→ 输入 111111 → 确认', '"所有数据都在你的手机里，不上传任何服务器。一键锁定，金额次数全变星号——手机借人、丢了都不怕。"'],
    ],
    [1.5, 5.5, 7]
)

# ═══════════════════════ 结尾 ═══════════════════════
doc.add_heading('结尾（2:50 - 3:00）', level=1)

add_table(
    ['时间', '画面操作', '旁白（建议）'],
    [
        ['2:50', '回到首页，全景停留 5 秒', '"青付安——买前有人帮你算账，买后有人帮你管卡，出事有人帮你留证。不用注册，数据就在你自己手里。"'],
    ],
    [1.5, 5.5, 7]
)

doc.add_page_break()

# ═══════════════════════ 设计说明 ═══════════════════════
doc.add_heading('设计说明', level=1)

doc.add_heading('四幕节奏设计', level=2)
add_table(
    ['幕', '时长', '核心画面', '对应评委关注点'],
    [
        ['痛点钩子', '20"', '到期提醒 + 首页全景', '用户洞察的准确性'],
        ['录卡→决策', '90"', '套餐录入流程 + 决策卡结论', '产品核心能力 + 规则引擎落地'],
        ['充卡差异化', '30"', '持仓卡（无限次）+ 核销打卡', '不是照搬记账软件的差异化设计'],
        ['证据夹 + 隐私', '30"', '材料清单 + 一键打包 + PIN 锁', '闭环价值 + 信任感'],
        ['收尾', '10"', '首页全景', '一句话总结定位'],
    ],
    [2, 1.5, 5, 5.5]
)

doc.add_heading('舍弃内容及原因', level=2)
add_table(
    ['舍弃', '原因'],
    [
        ['快速录入完整演示', '不是核心差异点，首页 ⊕ 弹窗已暗示存在'],
        ['风险报告独立页面', '决策卡已展示结论，报告页是信息重复'],
        ['首页场景标签切换', '无信息增量'],
        ['核销历史记录滚动', '打卡模式已传达差异点'],
        ['资产编辑/删除/管理', '管理类操作，无差异化价值'],
        ['"查看全部卡项"等导航按钮', '无信息增量'],
    ],
    [5, 9]
)

doc.add_heading('节奏控制要点', level=2)
add_para('• 各幕之间不留空档，边说边操作（提前准备剪贴板）', indent=0.5)
add_para('• 输入环节可后期 1.2× 加速，把时间留给决策卡展示', indent=0.5)
add_para('• 第一次出现决策卡时停顿 3 秒——这是全片最重要的静态画面', indent=0.5)
add_para('• 旁白用陈述句，不用反问句、感叹句——评委更接受冷静客观的语调', indent=0.5)

# ── 保存 ──
output = r'C:\Users\HUAWEI\Desktop\青付安演示视频录制大纲.docx'
doc.save(output)
print(f'Done: {output}')
